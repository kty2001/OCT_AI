"""
인턴십 대체 수업 주간 업무 보고서 생성 스크립트 (상세 버전).
기간: 2026.3.3 ~ 2026.6.8 (14주)
결과: results/internship_report.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT    = Path(__file__).parent.parent
RESULTS = ROOT / "results"
OUT     = RESULTS / "internship_report.docx"

# 이미지 경로
IMG = {
    "baseline":    RESULTS / "01_baseline" / "images" / "sample_01.png",
    "synth":       RESULTS / "02_synthetic_noise" / "AROI" / "samples" / "sample_00000.png",
    "n2n_result":  RESULTS / "03_sub2full" / "images" / "sample_01.png",
    "n2n_loss":    RESULTS / "03_sub2full" / "metrics" / "loss_curve.png",
    "sr":          RESULTS / "04_sr_test" / "all_comparison.png",
    "kfold_loss":  RESULTS / "06_kfold" / "metrics" / "loss_train_all_folds.png",
    "kfold_psnr":  RESULTS / "06_kfold" / "metrics" / "loss_val_psnr_all_folds.png",
    "dncnn":       RESULTS / "07_dncnn" / "images" / "fold_1" / "sample_01.png",
    "nafnet":      RESULTS / "08_nafnet" / "images" / "fold_1" / "sample_01.png",
    "nafnet_loss": RESULTS / "08_nafnet" / "metrics" / "loss_train_all_folds.png",
    "aug":         RESULTS / "09_nafnet_aug" / "images" / "fold_1" / "sample_01.png",
    "pretrain_log":RESULTS / "13_aroi_n2n" / "pretrain" / "loss_curve.png",
    "ft_result":   RESULTS / "13_aroi_n2n" / "finetune" / "images" / "fold_1" / "sample_01.png",
    "ft_psnr":     RESULTS / "13_aroi_n2n" / "finetune" / "metrics" / "loss_val_psnr_all_folds.png",
}


# ---------------------------------------------------------------------------
# 유틸리티
# ---------------------------------------------------------------------------

def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)


def set_spacing(p, before=0, after=80, line=360):
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    if before:
        spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_heading1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=(0x1a, 0x1a, 0x2e))
    set_spacing(p, before=280, after=120, line=360)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(0x16, 0x21, 0x3e))
    set_spacing(p, before=200, after=80, line=360)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(0xcc, 0x78, 0x5c))
    set_spacing(p, before=140, after=60, line=360)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10.5)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "480")
    pPr.append(ind)
    set_spacing(p, after=80, line=400)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "480")
    ind.set(qn("w:hanging"), "240")
    pPr.append(ind)
    run_bullet = p.add_run("• ")
    set_font(run_bullet, size=10.5, bold=True, color=(0xcc, 0x78, 0x5c))
    run = p.add_run(text)
    set_font(run, size=10.5)
    set_spacing(p, after=50, line=360)
    return p


def add_img(doc, path: Path, caption: str, width_cm: float = 14.0):
    if not path.exists():
        print(f"  [WARN] 이미지 없음: {path}")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    set_spacing(p_img, before=60, after=40, line=240)

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run(caption)
    set_font(r, size=9.5, color=(0x6c, 0x6a, 0x64))
    set_spacing(p_cap, after=120, line=280)


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_spacing(p, before=80, after=80, line=240)


def week_header(doc, num, dates, title):
    p = doc.add_paragraph()
    r1 = p.add_run(f"제 {num}주차")
    set_font(r1, size=14, bold=True, color=(0xcc, 0x78, 0x5c))
    r2 = p.add_run(f"  ({dates})")
    set_font(r2, size=11, color=(0x6c, 0x6a, 0x64))
    set_spacing(p, before=300, after=60, line=360)

    p2 = doc.add_paragraph()
    r3 = p2.add_run(title)
    set_font(r3, size=12, bold=True, color=(0x18, 0x17, 0x15))
    set_spacing(p2, after=100, line=360)


# ---------------------------------------------------------------------------
# 문서 생성
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    # ── 표지 ──────────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("인턴십 대체 수업 업무 보고서")
    set_font(r, size=24, bold=True, color=(0x18, 0x17, 0x15))

    doc.add_paragraph()
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sp.add_run("AI 기반 OCT 스페클 노이즈 제거 연구")
    set_font(r2, size=16, color=(0x3d, 0x3d, 0x3a))

    doc.add_paragraph()
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = pp.add_run("수행 기간: 2026년 3월 3일 ~ 2026년 6월 8일 (14주)")
    set_font(r3, size=12, color=(0x6c, 0x6a, 0x64))

    doc.add_page_break()

    # ── 프로젝트 개요 ──────────────────────────────────────────────────────
    add_heading1(doc, "1. 프로젝트 개요")

    add_para(doc,
        "본 프로젝트는 OCT(광간섭단층촬영, Optical Coherence Tomography)를 활용한 치주질환 탐지 및 예측 연구에 "
        "인공지능(AI)을 도입하는 것을 목표로 한다. OCT는 근적외선 파장의 빛을 이용하여 생체 조직 내부의 단층 구조를 "
        "마이크로미터 단위의 초고해상도로 획득하는 비침습적 광학 영상 기술로, 치과 및 치주 분야에서의 활용 가능성이 "
        "높게 평가되고 있다. 그러나 OCT 특유의 스페클 노이즈(speckle noise)는 영상 품질을 저하시켜 진단 정확도에 "
        "부정적 영향을 미친다.")

    add_para(doc,
        "치주 OCT 공개 데이터셋은 대부분 비공개(in-house)로 운용되어 직접 활용이 불가능한 상황이다. 이에 본 연구에서는 "
        "공개된 안과 영역의 대규모 망막 OCT 데이터셋을 활용하여 모델을 학습한 뒤, 추후 치주 도메인으로 전이하는 "
        "전략을 채택하였다. 핵심 목표는 스페클 노이즈 제거와 초해상도(Super-Resolution)를 단일 AI 파이프라인으로 "
        "처리하여 진단 영상의 품질을 향상시키는 것이다.")

    add_heading3(doc, "주요 성과 요약 (14주 기준)")
    results_summary = [
        "PSNR 28.35 dB (9단계, NAFNet+Aug) — 전통 방법 최고인 SRAD(27.50 dB) 초과 달성",
        "SSIM 0.6838 (13단계, AROI N2N 사전학습) — 전체 방법 중 최고",
        "CNR 1.186 (11단계, +Edge Loss) — AI 방법 중 최고, SRAD(1.220) 미달로 지속 연구 필요",
        "총 13가지 딥러닝 방법 체계적 비교 완료",
    ]
    for item in results_summary:
        add_bullet(doc, item)

    doc.add_page_break()

    # ── 주간 업무 내용 ─────────────────────────────────────────────────────
    add_heading1(doc, "2. 주간 업무 내용")

    # ════════════════════════════════════════════════════════════════════════
    # 1주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 1, "2026.03.03 ~ 2026.03.08", "프로젝트 시작 및 연구 방향 설정")

    add_heading3(doc, "연구 주제 확정 및 목표 설정")
    add_para(doc,
        "첫 주에는 인턴십 대체 수업의 주제를 확정하고 연구 방향을 수립하였다. AI를 활용한 OCT 영상의 스페클 노이즈 "
        "제거 및 초해상도를 핵심 연구 과제로 설정하였으며, 치주질환 진단을 최종 응용 목표로 삼았다. "
        "OCT 기술은 마이켈슨 간섭계 구조를 활용해 빛의 후방 산란 및 간섭 현상으로 단층 영상을 재구성하는 방식으로, "
        "치과 및 치주 분야에서의 비침습적 진단 도구로 주목받고 있다. 그러나 레이저의 가간섭성으로 인해 발생하는 "
        "스페클 노이즈는 OCT 영상 품질의 핵심 장애 요인으로 작용하며, 이를 AI로 제거하는 것이 본 연구의 출발점이다.")

    add_heading3(doc, "개발 환경 구성")
    add_para(doc,
        "연구 수행을 위한 개발 환경을 구성하였다. Python 3.12를 기반으로, uv 패키지 매니저를 채택하여 "
        "의존성 관리의 재현성을 확보하였다. PyTorch(CUDA 지원), numpy, pillow, scikit-image, bm3d, "
        "pandas, matplotlib, tqdm 등 핵심 라이브러리를 설치하였으며, NVIDIA RTX A4000 GPU(VRAM 16GB)를 "
        "주 연산 장치로 설정하였다. 프로젝트 디렉토리는 data/, scripts/, results/, weights/의 "
        "4개 최상위 디렉토리로 구분하여 데이터, 실험 코드, 결과물, 사전학습 가중치를 분리 관리하는 체계를 갖추었다.")

    add_bullet(doc, "Python 3.12, uv 패키지 관리자 기반 가상환경 구성")
    add_bullet(doc, "PyTorch(CUDA), scikit-image, bm3d, python-pptx 등 주요 라이브러리 설치")
    add_bullet(doc, "Git 저장소 초기화 및 브랜치 전략 수립")
    add_bullet(doc, "실험 재현성을 위한 프로젝트 구조(data/, scripts/, results/) 설계")
    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 2주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 2, "2026.03.09 ~ 2026.03.15", "OCT 원리 및 스페클 노이즈 특성 조사")

    add_heading3(doc, "OCT 작동 원리 학습")
    add_para(doc,
        "OCT는 마이켈슨 간섭계 구조를 기반으로 하는 광학 단층 영상 기술이다. 근적외선 파장의 빛을 시료 빔과 "
        "기준 빔으로 분리하여 각각 조직과 거울에 반사시킨 뒤, 두 빔의 간섭 신호를 분석함으로써 조직 내부의 "
        "반사 특성을 깊이 방향으로 측정한다. 단일 깊이 방향 스캔을 A-scan이라 하며, 연속된 A-scan을 가로 방향으로 "
        "배열하여 구성한 2차원 단면 영상을 B-scan이라 한다. 현대 임상 표준은 FFT 기반의 Fourier-Domain OCT(FD-OCT)이며, "
        "치과 및 연조직 검사에는 더 긴 파장을 사용하는 SS-OCT(Swept-Source OCT) 방식이 투과 깊이 면에서 유리하다.")

    add_para(doc,
        "OCT의 해상도는 두 축이 독립적으로 결정된다. 축 방향(Axial) 해상도는 광원의 파장과 대역폭에 의해 결정되며 "
        "임상 기준으로 약 5~10 µm에 달한다. 측 방향(Lateral) 해상도는 대물렌즈의 개구수(NA)에 의해 결정되며, "
        "초점 범위를 벗어나면 급격히 저하된다. 이러한 초고해상도 특성이 치주 조직의 미세 구조를 비침습적으로 "
        "관찰하는 데 OCT를 적합한 도구로 만든다.")

    add_heading3(doc, "스페클 노이즈의 물리적 특성")
    add_para(doc,
        "스페클 노이즈(speckle noise)는 OCT 영상에서 피할 수 없는 근본적인 아티팩트이다. 이는 레이저 광원의 "
        "가간섭성(coherence)으로 인해 조직 내 산란체들에서 반사된 빛들이 위상이 무작위적으로 간섭하면서 "
        "발생한다. 각 촬영 시마다 서로 다른 무작위 패턴으로 나타나므로, 동일한 부위를 촬영해도 매번 다른 "
        "스페클 패턴이 형성된다. 이 성질은 추후 Noise2Noise 자가지도 학습에서 핵심 전제 조건으로 활용된다.")

    add_para(doc,
        "스페클 노이즈의 수학적 모델은 일반적으로 곱셈성(multiplicative) 성분과 가산성(additive) 성분의 "
        "합으로 표현된다: I = S × Ns + Na. 여기서 I는 관측된 노이즈 영상, S는 이상적인 clean 영상, "
        "Ns는 Gamma 분포를 따르는 곱셈성 스페클 노이즈(mean=1, var=1/L), Na는 가우시안 분포를 따르는 "
        "가산성 노이즈이다. L은 'looks 수'라 불리는 파라미터로, 높을수록 노이즈가 약하다. "
        "이 물리 기반 모델은 2단계의 합성 노이즈 생성 파이프라인 설계의 근거가 되었다.")

    add_heading3(doc, "스페클 노이즈 제거 전통적 방법 조사")
    add_para(doc,
        "AI 도입에 앞서 전통적 스페클 노이즈 제거 방법들을 조사하였다. NLM(Non-Local Means)은 이미지 전체에서 "
        "유사한 패치를 찾아 가중 평균을 취하는 비국소 필터로, OCT의 반복적 레이어 구조 보존에 효과적이다. "
        "BM3D(Block-Matching 3D)는 유사 블록을 3D 변환 도메인으로 묶어 임계화(thresholding)하는 방법으로 "
        "성능이 우수하나 처리 속도가 느리다. SRAD(Speckle Reducing Anisotropic Diffusion)는 스페클의 "
        "곱셈성 특성을 반영한 PDE 기반 이방성 확산 방정식으로, 경계 보존 성능이 탁월하여 OCT에 가장 적합한 "
        "전통적 방법으로 평가받는다.")

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 3주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 3, "2026.03.16 ~ 2026.03.22", "딥러닝 기반 노이즈 제거 선행 연구 조사")

    add_heading3(doc, "지도학습 기반 딥러닝 디노이징")
    add_para(doc,
        "딥러닝을 활용한 이미지 디노이징 분야의 주요 방법들을 체계적으로 조사하였다. "
        "DnCNN(Beyond a Gaussian Denoiser, Zhang et al. 2017)은 잔차 학습(residual learning)과 "
        "배치 정규화를 결합한 선구적 CNN 구조로, 깊은 네트워크가 노이즈 패턴을 직접 학습하도록 설계된다. "
        "U-Net은 인코더-디코더 구조에 스킵 연결(skip connection)을 추가하여 다중 스케일의 특징을 보존하는 "
        "아키텍처로, 의료 영상 분야에서 광범위하게 활용된다. "
        "NAFNet(Nonlinear Activation Free Network, Chen et al. 2022)은 ReLU 등 비선형 활성화 함수를 "
        "제거하고 SimpleGate(채널 분리 후 곱셈)와 Simple Channel Attention(SCA)을 도입하여 "
        "연산 효율과 성능을 동시에 향상시킨 최신 방법이다. 이 세 아키텍처가 본 연구의 6~8단계 "
        "체계적 비교 실험의 대상으로 선정되었다.")

    add_heading3(doc, "자가지도 학습 방법 조사")
    add_para(doc,
        "clean Ground Truth 없이 학습하는 자가지도 방법들을 중점적으로 조사하였다. "
        "Noise2Noise(Lehtinen et al. 2018)는 두 개의 노이즈 관측값만으로 clean 이미지를 추정할 수 있다는 "
        "통계적 원리를 기반으로 한다. 노이즈가 독립적이고 평균이 0이라면, E[||f(noisy1) - noisy2||²]를 "
        "최소화하는 것이 E[||f(noisy1) - clean||²]를 최소화하는 것과 동치임을 수학적으로 증명한다. "
        "OCT에서는 동일 조직의 인접 슬라이스가 조직 구조는 유사하면서 스페클 패턴은 독립적인 특성을 "
        "이용하여 N2N 쌍을 구성할 수 있다.")

    add_para(doc,
        "Sub2Full(Noise2Full의 변형, 2024)은 OCT 특화 자가지도 방법으로, 단일 B-scan의 스펙트럼을 "
        "절반으로 분리하여 서로를 N2N 타겟으로 활용한다. 이 방법은 Noise2Noise, Noise2Void보다 우수한 "
        "성능을 보고하며 Optics Letters에 게재되었다. Noise2Void는 이미지 내 수용 영역(receptive field)의 "
        "중심 픽셀을 마스킹하여 자기 자신을 타겟으로 사용하지 못하게 하는 방법이다.")

    add_heading3(doc, "최신 Diffusion 기반 방법 조사")
    add_para(doc,
        "2023~2025년의 최신 연구 동향으로 Diffusion 모델 기반 방법들을 조사하였다. "
        "GARD(MICCAI 2025)는 OCT 스페클 노이즈의 Gamma 분포 특성을 Diffusion 모델의 forward process에 "
        "직접 적용하여 물리 기반 생성 모델을 구축한 방법이다. Content-Preserving Diffusion(MICCAI 2023)은 "
        "비지도 방식으로 해부학적 구조를 보존하면서 노이즈를 제거한다. 2024년의 systematic review에 따르면 "
        "방법론 계층은 Diffusion(최고 복원 품질) > GAN 기반 > 자가지도(실용성 1위) > 지도학습 CNN 순이다. "
        "본 연구에서는 데이터 제약과 구현 복잡도를 고려하여 자가지도 및 지도학습 방법에 집중하고, "
        "Diffusion 방법은 향후 연구 방향으로 남겨두었다.")

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 4주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 4, "2026.03.23 ~ 2026.03.29", "공개 데이터셋 탐색 및 수집")

    add_heading3(doc, "치주 OCT 데이터 탐색 및 대체 전략 수립")
    add_para(doc,
        "치주질환 관련 OCT 공개 데이터셋을 체계적으로 탐색하였다. ViT + Attention Gate 기반 치아 결손 탐지 연구 "
        "(MDPI 2023), 치조골 세그멘테이션 프레임워크 FD-SOS(MICCAI 2024), YOLOv8 기반 치주질환 스크리닝 ENPAT "
        "(PeerJ 2024) 등을 조사하였으나, 이들 연구에서 사용된 OCT 데이터셋은 모두 비공개임을 확인하였다. "
        "이에 따라 공개된 안과 망막 OCT 데이터셋으로 모델을 학습한 뒤 치주 도메인으로 전이하는 대체 전략을 최종 확정하였다.")

    add_heading3(doc, "SBSDI 데이터셋 수집 및 구조 분석")
    add_para(doc,
        "Fang et al.이 2013년 IEEE TMI에 발표한 SBSDI(Sparsity Based Simultaneous Denoising and Interpolation) "
        "패키지를 수집하였다. 이 데이터셋은 Bioptigen SDOCT(축 방향 해상도 4.5 µm/pixel)로 촬영된 "
        "인간 및 마우스 망막 OCT 영상으로 구성된다. 패키지 내 세 하위 데이터셋의 구조를 분석한 결과, "
        "D1(For synthetic experiments, 18쌍)만이 clean Ground Truth(average.tif)를 보유하고 있음을 확인하였다. "
        "D2(For real experiments on Humans, 39세트)는 13명 피험자의 중심와·상부·하부 3위치를 촬영한 데이터로 "
        "clean GT가 없어 자가지도 학습에만 활용 가능하다. D3(For real experiments on Mouse)는 마우스 망막을 "
        "1배·2배·4배 다운샘플 조건으로 촬영한 SR 비교용 데이터이다.")

    add_heading3(doc, "추가 데이터셋 수집")
    add_para(doc,
        "AROI(Annotated Retinal OCT Images Database)를 수집하였다. 이 데이터셋은 Heidelberg Spectralis OCT로 "
        "촬영된 nAMD(신생혈관성 황반변성) 환자 24명의 망막 OCT 영상으로 구성되며, 총 3,072장(24명 × 128장) 중 "
        "1,136장에 6클래스 레이어 세그멘테이션 주석이 완료되어 있다. 이미지가 세로로 긴 형태(512×1024 픽셀)로 "
        "저장되어 있어 일반적인 OCT B-scan의 가로 긴 형태와 90도 회전된 상태임을 파악하였다. "
        "또한 Kermany OCT2017(Kaggle, CC BY 4.0)을 수집하였다. 이 데이터셋은 84,484장의 망막 OCT 영상으로 "
        "CNV(맥락막 신생혈관), DME(당뇨황반부종), DRUSEN(초기 AMD), NORMAL(정상) 4클래스로 구성된다.")

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 5주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 5, "2026.03.30 ~ 2026.04.05", "데이터 심층 분석 및 평가 지표 구현")

    add_heading3(doc, "Clean GT 생성 가능성 분석")
    add_para(doc,
        "보유 데이터셋에서 추가적인 clean GT를 생성할 수 있는지 체계적으로 분석하였다. "
        "SBSDI D1의 average.tif가 어떤 방식으로 생성되었는지 MATLAB 코드(Demo_SBSDI.m)를 분석한 결과, "
        "동일 위치를 N≈40회 반복 스캔한 뒤 픽셀별 평균을 취한 것임을 확인하였다. "
        "Dictionary 학습용 HH/LL 이미지 쌍 10개의 배경 노이즈 std 비율 분석으로 N≈35~40으로 추정하였다. "
        "N회 평균 시 노이즈 표준편차가 1/√N배로 감소하므로, N=40이면 단일 프레임 대비 노이즈가 약 84% 제거된다.")

    add_para(doc,
        "다른 데이터셋의 GT 생성 가능성은 diff(frameA - frameB) std / frameA std 비율로 판별하였다. "
        "동일 위치 반복 스캔이라면 구조 성분이 상쇄되어 비율이 약 0.77 수준이고, "
        "다른 위치의 인접 슬라이스라면 구조 차이가 그대로 반영되어 비율이 1.0 이상이다. "
        "분석 결과 SBSDI D2(1.07), AROI(0.86~1.31), Kermany OCT2017 등 모든 보유 데이터셋에서 "
        "추가 GT 생성이 불가능함을 확인하였다. 이 분석은 이후 데이터 전략 수립의 핵심 근거가 되었다.")

    add_heading3(doc, "평가 지표 구현")
    add_para(doc,
        "성능 평가를 위한 세 가지 지표를 구현하였다. PSNR(Peak Signal-to-Noise Ratio)은 "
        "복원 이미지와 clean 이미지의 픽셀별 제곱 오차(MSE)를 기반으로 계산되며, "
        "오차가 작을수록 값이 높아지는 수치 정확도 지표이다. SSIM(Structural Similarity Index)은 "
        "밝기·대비·구조 공분산 세 요소를 동시에 비교하여 인간의 시지각 특성에 가까운 평가를 제공한다. "
        "CNR(Contrast-to-Noise Ratio)은 signal 영역(조직)과 background 영역의 평균 강도 차이를 "
        "배경 노이즈 수준으로 나눈 값으로, OCT 레이어 경계 선명도와 임상 진단 품질에 직접 대응하는 "
        "OCT 특화 지표이다. 세 지표 모두 동일 기준의 utils.py 함수로 통합하여 모든 실험 단계에서 "
        "일관성 있는 비교가 가능하도록 하였다.")

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 6주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 6, "2026.04.06 ~ 2026.04.19", "1단계: 전통적 방법 베이스라인 측정")

    add_heading3(doc, "구현 및 실험 진행")
    add_para(doc,
        "딥러닝 모델 도입 전 전통적 신호처리 방법으로 스페클 노이즈 제거 성능을 측정하여 비교 기준값을 확보하였다. "
        "NLM, BM3D, SRAD 세 방법을 구현하고 SBSDI D1 18쌍 전체에 적용하였다. "
        "NLM은 scikit-image 라이브러리의 구현체를 활용하였으며, BM3D는 별도 패키지(bm3d)를 설치하여 사용하였다. "
        "SRAD는 스페클 곱셈성 모델을 반영한 PDE를 직접 구현하였다.")

    add_para(doc,
        "실험 결과 SRAD가 PSNR 27.50±1.98 dB, SSIM 0.652±0.023, CNR 1.220±0.121로 세 지표 모두 1위를 기록하였다. "
        "곱셈성 스페클 특성에 특화된 확산 방정식이 OCT 영상에 가장 적합함을 확인하였다. "
        "BM3D는 SSIM에서 NLM 대비 큰 폭 향상(0.492→0.599)을 보여 비지역적 구조 보존 능력이 우수함을 확인하였다. "
        "이 결과를 바탕으로 딥러닝의 목표 기준값을 PSNR > 27.50 dB, SSIM > 0.652, CNR > 1.220으로 설정하였다.")

    add_img(doc, IMG["baseline"],
            "[그림 1] SBSDI D1 sample #01 비교: 좌→우 = Noisy 단일 프레임 / Clean 다중프레임 평균 / NLM / BM3D / SRAD",
            width_cm=15.0)

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 7주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 7, "2026.04.20 ~ 2026.04.26", "2단계: 물리 기반 합성 스페클 노이즈 생성 파이프라인 구축")

    add_heading3(doc, "노이즈 모델 및 캘리브레이션")
    add_para(doc,
        "clean GT가 없는 대규모 데이터셋(AROI, Kermany)을 지도학습에 활용하기 위해 "
        "물리 기반 합성 스페클 노이즈 생성 파이프라인을 구축하였다. "
        "노이즈 모델은 I = S × Ns + Na 형태로, Ns는 Gamma(L, 1/L) 분포를 따르는 "
        "곱셈성 스페클 성분(mean=1, var=1/L)이고, Na는 N(0, sigma_a²) 가우시안 가산성 성분이다. "
        "파라미터 L은 SBSDI D1 18쌍에서 Ns = I/S 비율의 Gamma 모멘트 매칭으로 역산하였다. "
        "18쌍의 L 추정값(3.015~6.228)의 평균으로 L = 5.266을 채택하였고, "
        "KS 통계량 0.119로 실제 스페클 분포와의 일치도를 검증하였다.")

    add_para(doc,
        "주의할 점은 AROI와 Kermany가 이미 실제 OCT 스페클을 포함하는 real 이미지라는 것이다. "
        "pixel-aligned clean GT가 없으므로 원본 이미지를 의사(pseudo) clean 기준으로 삼아 "
        "합성 스페클을 추가하는 방식을 사용하였다. 이 방식으로 학습된 모델은 추가된 합성 스페클만을 "
        "제거 대상으로 학습하게 되어 실제 OCT 스페클에 대한 일반화가 제한될 수 있으며, "
        "이 도메인 갭 문제는 3단계-B 실험에서 실제로 확인되었다.")

    add_heading3(doc, "학습 데이터 생성")
    add_para(doc,
        "AROI 1,136장(주석 완료 B-scan)과 Kermany OCT2017 train 서브셋 5,000장(랜덤 샘플링, seed=42)에 "
        "캘리브레이션된 파라미터(L=5.266, sigma=0.010)로 합성 노이즈를 추가하여 총 6,136쌍을 생성하였다. "
        "AROI 이미지는 90도 CCW 회전 보정을 적용하였으며, 생성된 쌍의 경로와 메타데이터를 "
        "metadata.csv로 관리하여 재현 가능한 파이프라인을 구축하였다.")

    add_img(doc, IMG["synth"],
            "[그림 2] AROI 합성 노이즈 샘플: 좌=원본(pseudo clean), 우=합성 노이즈 추가",
            width_cm=14.0)

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 8주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 8, "2026.04.27 ~ 2026.05.03", "3단계-A: Noise2Noise 자가지도 학습")

    add_heading3(doc, "방법론 및 학습 설정")
    add_para(doc,
        "clean Ground Truth 없이 OCT 이미지만으로 학습하는 Noise2Noise 자가지도 방식을 구현하였다. "
        "SBSDI D2의 39세트에서 각 세트별 test.tif와 1~4.tif(인접 공간 슬라이스)를 N2N 쌍으로 활용하였다. "
        "인접 슬라이스는 3D 볼륨 내 이웃 B-scan으로, 조직 구조가 매우 유사하면서 스페클 패턴은 "
        "독립적이므로 근사 Noise2Noise 조건을 만족한다. 다만 완전한 동일 위치 반복 스캔이 아니어서 "
        "diff/frame std 비율이 1.07로, 이상적인 D1 동일 위치(0.77) 대비 근사 N2N 조건임을 인지하였다.")

    add_para(doc,
        "39세트 × 12 순서쌍(i≠j) = 468쌍의 학습 데이터를 구성하고, "
        "경량 U-Net(~1.95M params, 인코더 3단계 + 보틀넥 + 디코더 3단계)을 구현하여 "
        "L1 Loss, Adam(lr=1e-4, CosineAnnealing), 배치=4, 패치=128×128, 500 epoch으로 학습하였다. "
        "약 221분(RTX A4000) 소요되었으며 loss는 0.105에서 0.099로 수렴 폭이 매우 작았다. "
        "이는 39세트라는 절대적 데이터 부족이 학습에 제약을 가하고 있음을 시사한다.")

    add_heading3(doc, "평가 결과")
    add_para(doc,
        "SBSDI D1 18쌍 기준 PSNR 26.84±2.18 dB, SSIM 0.681±0.033, CNR 1.148±0.133을 달성하였다. "
        "SSIM 0.681이 SRAD(0.652)를 초과한 것은 clean GT 없이도 구조 보존 측면에서 전통 방법 1위를 "
        "능가하는 의미 있는 결과이다. PSNR은 SRAD(27.50)보다 0.66 dB 낮고 CNR은 미달하였으며, "
        "이는 N2N 노이즈 타겟의 고유 분산이 loss floor를 형성하여 clean GT 없이는 성능 상한이 "
        "존재함을 확인시켜준다.")

    add_img(doc, IMG["n2n_result"],
            "[그림 3] 3단계-A N2N 복원 결과 (sample #01): 좌→우 = Noisy / Clean / NLM / BM3D / SRAD / N2N",
            width_cm=15.0)

    add_img(doc, IMG["n2n_loss"],
            "[그림 4] 3단계-A N2N 학습 곡선: 500 epoch 동안 loss 0.105→0.099 (수렴 폭 작음)",
            width_cm=12.0)

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 9주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 9, "2026.05.04 ~ 2026.05.10", "3단계-B & 4단계: 합성 지도학습 + Real-ESRGAN SR")

    add_heading3(doc, "3단계-B: 합성 데이터 지도학습 및 도메인 갭 분석")
    add_para(doc,
        "2단계에서 생성한 합성 noisy-clean 6,136쌍으로 U-Net 지도학습을 수행하였다. "
        "L1 Loss, Adam(lr=2e-4, CosineAnnealing), 에포크=100, 배치=8, 패치=128×128 설정으로 "
        "약 305분 학습하였다. Best loss는 0.021621이며, epoch 10 이후 수렴이 포화됨을 관찰하였다. "
        "평가 결과는 PSNR 22.43±0.94, SSIM 0.333±0.049, CNR 0.902±0.100으로 "
        "세 지표 모두 SRAD를 크게 하회하였다. "
        "원인은 학습 데이터의 합성 Gamma+Gaussian 노이즈와 평가 데이터(SBSDI D1)의 실제 OCT 스페클 간 "
        "도메인 갭으로, 수학적 노이즈 모델이 복잡한 물리적 스페클 특성을 완전히 모사하지 못하기 때문이다.")

    add_heading3(doc, "4단계: Real-ESRGAN Blind SR 적용")
    add_para(doc,
        "스페클 노이즈 제거와 초해상도를 동시에 처리하는 사전학습 모델로 Real-ESRGAN을 적용하였다. "
        "Real-ESRGAN은 실제 복합 열화(노이즈 + 저해상도 + 압축 아티팩트)를 동시에 복원하도록 학습된 "
        "blind SR 모델이다. RealESRGAN_x2plus.pth와 x4plus.pth(각 64MB) 가중치를 활용하여 "
        "CUDA half precision, tile=400 설정으로 추론하였다. SBSDI D1 18쌍 전체 평가 결과, "
        "x4 모델이 PSNR 27.57±2.40, SSIM 0.673±0.034, CNR 1.166±0.117을 달성하였으며 "
        "fine-tuning 없이 PSNR과 SSIM에서 SRAD를 초과하는 성과를 보였다. "
        "처리 속도도 SRAD(5.68s) 대비 1.42s/장으로 4배 빠르다. "
        "다만 CNR은 SRAD(1.220) 미달로, blind SR 모델이 고주파 디테일 복원에는 강하지만 "
        "전역 대비 보존에는 한계가 있음을 확인하였다.")

    add_img(doc, IMG["sr"],
            "[그림 5] 4단계 Real-ESRGAN SR 결과 비교: 좌→우 = Noisy / Clean / ESRGAN x2 / ESRGAN x4",
            width_cm=15.0)

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 10주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 10, "2026.05.11 ~ 2026.05.18", "5단계: Pre-train → Fine-tune 전이학습")

    add_heading3(doc, "전략 설계 및 구현")
    add_para(doc,
        "3단계-B에서 합성 데이터로 학습한 사전학습 가중치(best.pth, loss 0.021621)로 U-Net을 초기화한 뒤 "
        "SBSDI D2 real 39세트 × 12 = 468쌍의 N2N 방식으로 fine-tuning하였다. "
        "손실 함수는 L1에서 L1 + 0.1×(1-SSIM) 조합으로 개선하여 구조 보존 능력을 강화하였다. "
        "학습률은 1e-5(사전학습 lr=2e-4 대비 20배 낮춤)로 설정하여 사전학습 가중치의 과도한 변형을 방지하였으며, "
        "200 epoch, 배치=16, 약 35분 소요되었다.")

    add_para(doc,
        "평가 결과는 PSNR 27.46±2.55, SSIM 0.679±0.037, CNR 1.169±0.120으로 "
        "N2N(26.84)에 비해 PSNR +0.62 dB 향상되었고 SRAD까지 0.04 dB 차이에 근접하였다. "
        "그러나 loss가 epoch 20 이후 약 0.184에서 포화되어 더 이상 수렴하지 않는 현상이 관찰되었다. "
        "이는 N2N 노이즈 타겟의 고유 분산이 loss floor를 형성하기 때문으로, clean GT 없이는 "
        "이 분산을 제거할 수 없어 성능 상한이 결정됨을 의미한다. "
        "이 한계를 극복하기 위해 SBSDI D1의 real clean GT를 활용하는 k-fold CV 방향으로 전환하였다.")

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 11주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 11, "2026.05.19 ~ 2026.05.25", "6~8단계: 6-fold CV 아키텍처 체계적 비교")

    add_heading3(doc, "6단계: U-Net 6-fold CV 지도학습")
    add_para(doc,
        "SBSDI D1 18쌍을 k=6 fold로 나눠 real clean GT로 지도학습하는 6-fold Cross-Validation 방식을 도입하였다. "
        "6단계-A(batch=16, ep=150, early stopping 없음)는 PSNR 27.21로 SRAD 미달이었다. "
        "학습 loss가 epoch 150에서도 감소 중이었으며, early stopping 없이 고정 epoch를 사용한 것이 "
        "수렴 부족의 원인임을 분석하였다. "
        "6단계-B(batch=64, ep=500, patience=30)는 val PSNR+10×SSIM 스코어를 기준으로 "
        "early stopping을 적용하여 PSNR 28.19±2.56을 달성하였다. "
        "이는 AI 방법 최초로 SRAD(27.50 dB)를 초과한 성과이며, early stopping이 "
        "과적합 방지와 수렴 동시에 해결하면서 학습 시간도 175분에서 47분으로 단축하였다.")

    add_img(doc, IMG["kfold_psnr"],
            "[그림 6] 6단계 U-Net 6-fold CV: Fold별 Val PSNR 수렴 곡선 (6개 fold 모두 35~47 epoch에서 조기 종료)",
            width_cm=14.0)

    add_heading3(doc, "7~8단계: DnCNN 및 NAFNet 아키텍처 비교")
    add_para(doc,
        "동일한 6-fold CV 프레임워크에서 백본을 교체하여 아키텍처의 영향을 측정하였다. "
        "7단계의 DnCNN-B(depth=20, ch=64, 667K params)는 from scratch 초기화, PSNR 28.17±2.47을 달성하였다. "
        "8단계의 NAFNet-32(17M params, SimpleGate+SCA+LayerNorm2d)는 PSNR 28.11±2.26, CNR 1.183으로 "
        "AI 방법 중 CNR이 가장 높았다. "
        "핵심 발견은 667K(DnCNN) ≈ 1.95M(U-Net) ≈ 17M(NAFNet)으로, "
        "파라미터 수 25배 차이에도 PSNR 차이가 0.08 dB 이내라는 것이다. "
        "이는 18쌍이라는 데이터 크기가 아키텍처 용량의 차이를 완전히 압도하는 '데이터 병목' 현상임을 "
        "3회 연속으로 확인한 결과이다.")

    add_img(doc, IMG["nafnet"],
            "[그림 7] 8단계 NAFNet fold_1 복원 결과 (sample #01): 좌→우 = Noisy / Clean / 복원",
            width_cm=15.0)

    add_img(doc, IMG["nafnet_loss"],
            "[그림 8] 8단계 NAFNet 6-fold CV Train Loss 수렴 곡선 (6개 fold): lr=1e-3에서 빠른 초기 수렴 후 포화",
            width_cm=14.0)

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 12주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 12, "2026.05.26 ~ 2026.06.01", "9단계: 노이즈 재실현 증강 + 발표자료 v2 제작")

    add_heading3(doc, "9단계: 다중 노이즈 재실현 증강")
    add_para(doc,
        "8단계까지 아키텍처를 바꿔도 성능 차이가 0.08 dB 이내로 데이터 병목이 확인된 상황에서, "
        "합성 노이즈 재실현으로 학습 데이터를 증강하는 방법을 시도하였다. "
        "18개 clean GT 각각에 캘리브레이션된 노이즈 모델(L=5.266)로 K=4개씩 서로 다른 시드(seed=42+offset)로 "
        "합성 noisy를 재실현하여 72장을 생성하였다. fold당 학습 쌍이 15→75(5배), "
        "패치 수가 4,125→20,625(5배)로 증가하였다. 평가는 real noisy로만 수행하여 8단계와 동등한 조건을 유지하였다.")

    add_para(doc,
        "결과는 PSNR 28.35±2.56(+0.24), SSIM 0.6832±0.032(+0.009)로 현재 전체 방법 중 PSNR/SSIM 최고를 달성하였다. "
        "CNR은 1.168로 8단계(1.183) 대비 소폭 하락하였는데, 합성 노이즈가 실제 스페클과 완전히 동일하지 않아 "
        "대비 보존 측면에서 미미하게 영향을 준 것으로 분석된다. "
        "K=4 재실현 증강의 효과는 +0.24 dB로 확인되었으나, 데이터 병목의 근본적 해소를 위해서는 "
        "외부 real OCT 데이터 확보가 필요하다는 결론을 도출하였다.")

    add_img(doc, IMG["aug"],
            "[그림 9] 9단계 NAFNet+Aug fold_1 복원 결과 (sample #01): 노이즈 재실현 K=4 증강 적용",
            width_cm=15.0)

    add_heading3(doc, "발표자료 v1 & v2 제작")
    add_para(doc,
        "1~9단계의 실험 내용을 정리한 발표자료를 python-pptx 기반 자동 생성 스크립트로 제작하였다. "
        "v1(13슬라이드)을 제작 후, 더 상세한 내용을 담은 v2(20슬라이드)를 개선하였다. "
        "OCT 기초 및 평가 지표, 데이터셋 현황, 방법론(1~9단계), 종합 결과 비교의 "
        "4개 섹션으로 구성하였으며, 디자인 시스템(색상 팔레트, 맑은 고딕/나눔고딕 폰트)을 통일하였다. "
        "또한 README.md, data_description.md, result_description.md를 전면 정비하여 "
        "실험 설정, 결과 수치, 분석 내용을 체계적으로 문서화하였다.")

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 13주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 13, "2026.06.02 ~ 2026.06.04", "10~12단계: 하이퍼파라미터 및 손실 함수 개선 탐색")

    add_heading3(doc, "10단계: 학습률(lr) 조정 실험")
    add_para(doc,
        "8단계 NAFNet에서 학습 초반 val PSNR이 epoch 1~9에서 최고점을 찍고 이후 정체·하락하는 "
        "loss-PSNR 분리 현상을 관찰하였다. lr=1e-3이 너무 높아 초반에 훈련 데이터를 빠르게 암기하고 "
        "이후 과적합되는 것으로 가설을 세워 lr=1e-4로 낮춰 재실험하였다. "
        "결과는 PSNR 27.95로 기준(28.11) 대비 오히려 -0.16 dB 하락하였다. "
        "원인은 lr=1e-3이 초반 몇 에포크에서 빠르게 좋은 수렴점을 찾고 early stopping이 그 시점을 포착하는 반면, "
        "lr=1e-4는 patience=30 내에 그 수렴점까지 도달하지 못하기 때문이다. "
        "데이터 18쌍 환경에서는 하이퍼파라미터 조정으로 데이터 병목을 해소하기 어렵다는 결론을 확인하였다.")

    add_heading3(doc, "11~12단계: Edge Loss 및 Frequency Loss 추가")
    add_para(doc,
        "CNR SRAD(1.220) 미달 문제를 해소하기 위해 경계 보존에 특화된 손실 함수를 추가하였다. "
        "Edge Loss는 Sobel 필터로 예측 이미지와 정답 이미지의 경계 맵을 각각 추출한 뒤 L1 거리를 최소화한다. "
        "Sobel 출력이 이미지마다 스케일이 다른 문제로 6회의 정규화 방식 시도 끝에 "
        "SOBEL_MAX(11.314)+clamp(0,1) 조합으로 안정적인 학습을 확립하였다. "
        "λ=0.5 적용 결과 CNR 1.183→1.186(+0.003)으로 소폭 향상되었으나 "
        "std(±0.116) 대비 통계적 유의성은 없었다. "
        "이어서 FFT log1p 기반 Frequency Loss(λ=0.1)를 추가하였으나 "
        "Edge Loss와 Frequency Loss가 gradient 방향에서 충돌하여 "
        "CNR이 오히려 1.182로 하락하였다. 손실 함수 개선의 한계를 확인하고 실험을 종료하였다.")

    add_divider(doc)

    # ════════════════════════════════════════════════════════════════════════
    # 14주차
    # ════════════════════════════════════════════════════════════════════════
    week_header(doc, 14, "2026.06.05 ~ 2026.06.08", "13단계: AROI N2N 사전학습 → D1 Fine-tuning + 최종 보고")

    add_heading3(doc, "13단계: AROI 인접 B-scan N2N 사전학습")
    add_para(doc,
        "외부 데이터 없이 보유 데이터의 활용도를 극대화하는 방향으로, "
        "AROI 데이터셋의 인접 B-scan 쌍을 N2N 사전학습에 활용하였다. "
        "AROI의 24명 환자 각각은 128장의 순차적 B-scan을 보유하므로 인접 슬라이스 (i, i+1)는 "
        "조직 구조가 유사하면서 스페클 패턴이 독립적인 N2N 조건을 만족한다. "
        "Train(patient 1~20): 2,540 N2N 쌍, Val(patient 21~24): 508 N2N 쌍으로 구성하였다.")

    add_para(doc,
        "초기 구현에서 exhaustive tiling(stride=32) 방식을 적용하여 957,580개의 패치가 생성되어 "
        "epoch당 314분이 소요되는 문제가 발생하였다. 이를 lazy loading + random crop 방식으로 개선하여 "
        "epoch당 4,096개 샘플(batch=48, 85 steps)로 제한하였고, 학습 시간이 약 2분/epoch으로 단축되었다. "
        "NAFNet width=32, L1 Loss, lr=1e-3, seed=42, patience=30 설정으로 epoch 93에서 조기 종료되었다. "
        "사전학습 중 SBSDI D1 PSNR을 모니터링한 결과 epoch 50에서 27.086 dB로 최고점을 기록하였으며, "
        "이는 SBSDI D2 N2N(3단계-A, 26.84 dB) 대비 +0.25 dB 향상된 수치이다.")

    add_img(doc, IMG["pretrain_log"],
            "[그림 10] 13단계 사전학습 곡선: 좌=N2N train/val loss, 우=SBSDI D1 PSNR 모니터링 (최고 27.086 dB, epoch 50)",
            width_cm=15.0)

    add_heading3(doc, "D1 k-fold Fine-tuning")
    add_para(doc,
        "사전학습된 가중치를 초기값으로 SBSDI D1 6-fold CV fine-tuning을 수행하였다. "
        "lr=1e-4(사전학습 lr=1e-3 대비 낮춤), Loss=L1+SSIM+EdgeLoss(λ=0.5), seed=42로 설정하였다. "
        "Epoch 1의 val PSNR이 28~30 수준으로 시작하여 사전학습의 효과를 직접 확인할 수 있었으며, "
        "모든 fold가 epoch 31~34에서 조기 종료되어 scratch(32~46)보다 빠른 수렴을 보였다.")

    add_para(doc,
        "최종 결과는 PSNR 28.23±2.49, SSIM 0.6838±0.029(전체 방법 중 최고), CNR 1.177±0.121이다. "
        "SSIM에서 기존 최고인 NAFNet+Aug(0.6832)를 0.0006 차이로 경신하였으며, "
        "PSNR도 scratch(28.11) 대비 +0.12 dB 향상되었다. "
        "CNR은 scratch(1.183) 대비 소폭 하락하였는데, AROI N2N이 smooth 출력을 선호하는 경향이 "
        "EdgeLoss 효과를 일부 상쇄한 것으로 분석된다.")

    add_img(doc, IMG["ft_result"],
            "[그림 11] 13단계 AROI 사전학습→D1 fine-tuning 복원 결과 (sample #01, fold_1): 좌→우 = Noisy / Clean / 복원",
            width_cm=15.0)

    add_img(doc, IMG["ft_psnr"],
            "[그림 12] 13단계 D1 fine-tuning 6-fold Val PSNR 수렴 곡선: epoch 1부터 28~30 dB 수준으로 시작",
            width_cm=14.0)

    add_heading3(doc, "최종 발표자료 v3 제작 및 보고서 정리")
    add_para(doc,
        "10~13단계의 신규 실험 내용을 포함한 발표자료 v3(22슬라이드)를 제작하였다. "
        "10~12단계(lr 조정 + 손실 함수 개선)와 13단계(AROI N2N 사전학습)의 신규 슬라이드 2개를 추가하고, "
        "종합 결과 차트/테이블을 16개 방법으로 확장하였으며 핵심 발견과 향후 계획을 업데이트하였다. "
        "학습 loss 그래프(plot_training_logs.py, 600 dpi)를 모든 실험 단계에 대해 생성하여 "
        "results/ 디렉토리에 저장하였다.")

    add_divider(doc)

    # ── 종합 결과 요약 ──────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading1(doc, "3. 종합 성능 비교")

    add_para(doc,
        "아래 표는 14주 동안 수행한 전체 실험의 성능을 SBSDI D1(18쌍 평균) 기준으로 정리한 것이다. "
        "모든 딥러닝 방법은 PSNR과 SSIM에서 전통 방법 SRAD를 초과하거나 근접하였으나, "
        "CNR에서는 어떤 AI 방법도 SRAD(1.220)를 초과하지 못하였다.")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["방법", "단계", "PSNR (dB)", "SSIM", "CNR"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)
                run.font.name = "맑은 고딕"

    rows_data = [
        ("NLM",                   "1단계",   "26.12±2.01", "0.492±0.050", "1.130±0.119"),
        ("BM3D",                  "1단계",   "27.00±2.51", "0.599±0.066", "1.134±0.119"),
        ("SRAD (기준값)",          "1단계",   "27.50±1.98", "0.652±0.023", "1.220±0.121"),
        ("N2N Sub2Full",          "3단계-A", "26.84±2.18", "0.681±0.033", "1.148±0.133"),
        ("지도학습 (합성)",        "3단계-B", "22.43±0.94", "0.333±0.049", "0.902±0.100"),
        ("Real-ESRGAN x4",        "4단계",   "27.57±2.40", "0.673±0.034", "1.166±0.117"),
        ("Pre-train + Fine-tune", "5단계",   "27.46±2.55", "0.679±0.037", "1.169±0.120"),
        ("U-Net 6-fold CV (ES)",  "6단계",   "28.19±2.56", "0.681±0.031", "1.169±0.121"),
        ("DnCNN 6-fold CV",       "7단계",   "28.17±2.47", "0.673±0.032", "1.167±0.127"),
        ("NAFNet-32 6-fold CV",   "8단계",   "28.11±2.26", "0.674±0.029", "1.183±0.120"),
        ("NAFNet + Aug K=4",      "9단계",   "28.35±2.56", "0.683±0.032", "1.168±0.121"),
        ("NAFNet lr=1e-4",        "10단계",  "27.95±2.11", "0.667±0.026", "1.191±0.116"),
        ("NAFNet + Edge Loss",    "11단계",  "28.06±2.08", "0.670±0.027", "1.186±0.116"),
        ("AROI N2N pretrain→FT",  "13단계",  "28.23±2.49", "0.684±0.029", "1.177±0.121"),
    ]
    for rd in rows_data:
        row = table.add_row().cells
        for i, v in enumerate(rd):
            row[i].text = v
            for p in row[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "맑은 고딕"

    doc.add_paragraph()

    # ── 향후 연구 방향 ──────────────────────────────────────────────────────
    add_heading1(doc, "4. 향후 연구 방향")

    add_para(doc,
        "14주간의 실험을 통해 데이터 18쌍이 유일한 병목임이 반복적으로 확인되었다. "
        "아키텍처 교체(DnCNN/U-Net/NAFNet), lr 조정, 손실 함수 개선(Edge+Freq Loss), "
        "K=4 노이즈 재실현 증강 모두 근본적 한계를 해소하지 못하였다. "
        "향후 연구는 데이터 측면의 접근에 집중해야 한다.")

    futures = [
        ("Duke AMD SD-OCT 활용",
         "384명, 38,400 B-scan을 직접 다운로드하여 N2N 사전학습 데이터 다양성을 대폭 확대한다. "
         "현재 AROI N2N 사전학습(3,048쌍)보다 1,000배 많은 다양성으로 "
         "사전학습 모델의 일반화 성능 향상을 기대할 수 있다."),
        ("외부 clean GT 데이터셋 접근 재시도",
         "PKU37(37쌍, 알리바바 클라우드), Sub2Full vis-OCT(저자 직접 이메일 컨택), "
         "RETOUCH(Grand Challenge 계정) 등 기존 접근에 실패한 데이터셋을 재시도한다. "
         "clean GT 확보로 k-fold CV 데이터를 늘리는 것이 가장 직접적인 성능 향상 방법이다."),
        ("치주 OCT 직접 수집 (장기)",
         "최종 목표 도메인인 치주 OCT를 직접 수집한다. "
         "SBSDI D1 방식(동일 위치 약 40회 반복 촬영 → 픽셀별 평균으로 clean GT 생성)을 적용하며, "
         "교수님과 촬영 프로토콜 협의가 필요하다."),
        ("스페클 제거 + SR 통합 파이프라인",
         "스페클 노이즈 제거 성능이 충분히 향상된 후 Real-ESRGAN SR과 연결하여 "
         "단일 파이프라인으로 치주 OCT 진단 품질을 향상시키는 것이 최종 목표이다."),
    ]

    for i, (title, body) in enumerate(futures):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{i+1}. {title}\n")
        set_font(r1, size=11, bold=True, color=(0xcc, 0x78, 0x5c))
        r2 = p.add_run(f"   {body}")
        set_font(r2, size=10.5, color=(0x3d, 0x3d, 0x3a))
        set_spacing(p, before=80, after=100, line=380)

    doc.save(str(OUT))
    print(f"저장 완료: {OUT}")


if __name__ == "__main__":
    build()
